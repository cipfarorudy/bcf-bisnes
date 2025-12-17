from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import datetime
import os
import textwrap

base = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(base, exist_ok=True)

# 1) HTML sales page
html = """<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>Back-office + Financement + Acquisition | Abonnement Pro</title>
  <style>
    :root{
      --bg:#0b1220; --card:#101a2e; --muted:#93a4c7; --text:#eaf0ff;
      --accent:#5eead4; --accent2:#60a5fa; --danger:#fb7185;
      --border:rgba(255,255,255,.10);
      --shadow: 0 18px 60px rgba(0,0,0,.35);
      --radius:18px;
      --max:1100px;
      --font: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial, "Apple Color Emoji","Segoe UI Emoji";
    }
    *{box-sizing:border-box}
    body{margin:0;background:radial-gradient(1200px 600px at 70% 0%, rgba(96,165,250,.18), transparent 60%),
                     radial-gradient(900px 500px at 20% 10%, rgba(94,234,212,.14), transparent 55%),
                     var(--bg);
         color:var(--text); font-family:var(--font); line-height:1.55}
    a{color:inherit}
    .wrap{max-width:var(--max); margin:0 auto; padding:28px 18px 80px}
    header{display:flex; align-items:center; justify-content:space-between; gap:14px; padding:10px 0 18px}
    .brand{display:flex; gap:10px; align-items:center}
    .logo{width:38px; height:38px; border-radius:12px; background:linear-gradient(135deg,var(--accent),var(--accent2));
          box-shadow:0 10px 30px rgba(96,165,250,.25)}
    .brand h1{font-size:16px; margin:0}
    .brand small{display:block; color:var(--muted); margin-top:2px}
    .topcta{display:flex; gap:10px; align-items:center; flex-wrap:wrap}
    .btn{border:1px solid var(--border); background:rgba(255,255,255,.04); color:var(--text);
         padding:10px 14px; border-radius:12px; text-decoration:none; display:inline-flex; gap:10px; align-items:center;
         box-shadow:0 10px 28px rgba(0,0,0,.25)}
    .btn.primary{background:linear-gradient(135deg, rgba(94,234,212,.95), rgba(96,165,250,.95));
                 color:#051022; border-color:transparent; font-weight:700}
    .btn:hover{transform:translateY(-1px)}
    .hero{display:grid; grid-template-columns: 1.2fr .8fr; gap:18px; margin-top:10px}
    @media(max-width:900px){.hero{grid-template-columns:1fr}}
    .card{background:linear-gradient(180deg, rgba(255,255,255,.05), rgba(255,255,255,.03));
          border:1px solid var(--border); border-radius:var(--radius); padding:18px; box-shadow:var(--shadow)}
    .kicker{color:var(--accent); font-weight:700; letter-spacing:.2px}
    .title{font-size:38px; margin:8px 0 10px; line-height:1.15}
    @media(max-width:540px){.title{font-size:30px}}
    .subtitle{color:var(--muted); font-size:16px; margin:0 0 16px}
    .bullets{display:grid; gap:10px; margin:16px 0 0; padding:0; list-style:none}
    .bullets li{display:flex; gap:10px; align-items:flex-start; color:#dbe6ff}
    .dot{width:10px; height:10px; border-radius:999px; margin-top:7px;
         background:linear-gradient(135deg,var(--accent),var(--accent2)); flex:0 0 10px}
    .grid3{display:grid; grid-template-columns:repeat(3,1fr); gap:14px; margin-top:16px}
    @media(max-width:900px){.grid3{grid-template-columns:1fr}}
    .chip{display:inline-flex; gap:8px; align-items:center; padding:6px 10px; border-radius:999px;
          border:1px solid var(--border); background:rgba(255,255,255,.03); color:var(--muted); font-size:13px}
    .price{display:flex; align-items:baseline; gap:10px; margin-top:10px}
    .price b{font-size:34px}
    .price span{color:var(--muted)}
    .section{margin-top:18px}
    .section h2{margin:0 0 10px; font-size:20px}
    .muted{color:var(--muted)}
    .form{display:grid; gap:10px; margin-top:12px}
    label{font-size:13px; color:var(--muted)}
    input, select, textarea{
      width:100%; padding:12px 12px; border-radius:12px; border:1px solid var(--border);
      background:rgba(0,0,0,.25); color:var(--text); outline:none
    }
    textarea{min-height:90px; resize:vertical}
    .twocol{display:grid; grid-template-columns:1fr 1fr; gap:10px}
    @media(max-width:700px){.twocol{grid-template-columns:1fr}}
    .note{font-size:13px; color:var(--muted)}
    .ok{color:var(--accent)}
    footer{margin-top:22px; color:var(--muted); font-size:13px}
    .faq details{border:1px solid var(--border); border-radius:14px; padding:12px 12px; background:rgba(255,255,255,.03)}
    .faq summary{cursor:pointer; font-weight:700}
    .faq p{margin:10px 0 0}
  </style>
</head>
<body>
  <div class="wrap">
    <header>
      <div class="brand">
        <div class="logo" aria-hidden="true"></div>
        <div>
          <h1>Abonnement Pro — Back-office + Financement + Acquisition</h1>
          <small>Externalisation opérationnelle pour structures en croissance</small>
        </div>
      </div>
      <div class="topcta">
        <a class="btn" href="#offres">Voir les offres</a>
        <a class="btn primary" href="#diagnostic">Demander un diagnostic</a>
      </div>
    </header>

    <div class="hero">
      <div class="card">
        <div class="kicker">Objectif : structurer, financer, convertir</div>
        <div class="title">On gère votre back-office, vos financements et vos prospects, pendant que vous développez.</div>
        <p class="subtitle">Un service unique qui combine centre d'appels, ingénierie de financement, infrastructure digitale et tunnel de vente — avec reporting et pilotage.</p>

        <div>
          <span class="chip">Call center & qualification</span>
          <span class="chip">Financement & dossiers</span>
          <span class="chip">Nom de domaine & mails pro</span>
          <span class="chip">CRM & tunnel de vente</span>
        </div>

        <ul class="bullets">
          <li><span class="dot"></span><div><b>Plus de temps</b> : vous déléguez l'administratif et la relation entrante.</div></li>
          <li><span class="dot"></span><div><b>Plus de cash</b> : vous activez les bons dispositifs de financement (selon éligibilité).</div></li>
          <li><span class="dot"></span><div><b>Plus de clients</b> : un tunnel clair, une base prospects, des relances structurées.</div></li>
        </ul>

        <div class="section">
          <a class="btn primary" href="#diagnostic">Démarrer par un diagnostic (gratuit)</a>
          <div class="note">Réponse sous 24–48h ouvrées. Mise en place rapide après validation.</div>
        </div>
      </div>

      <div class="card" id="offres">
        <h2>Offres</h2>
        <p class="muted">Choisissez la formule adaptée. Passage PRO → PREMIUM possible.</p>

        <div class="section">
          <div class="kicker">PRO</div>
          <div class="price"><b>500€</b><span>/ mois</span></div>
          <ul class="bullets">
            <li><span class="dot"></span><div>Call center mutualisé + qualification standard</div></li>
            <li><span class="dot"></span><div>Diagnostic financement + orientation + check-list dossier</div></li>
            <li><span class="dot"></span><div>Nom de domaine + 5 mails pro</div></li>
            <li><span class="dot"></span><div>Tunnel standard + base prospects (CRM)</div></li>
          </ul>
        </div>

        <div class="section" style="margin-top:14px">
          <div class="kicker">PREMIUM</div>
          <div class="price"><b>1 000€</b><span>/ mois</span></div>
          <ul class="bullets">
            <li><span class="dot"></span><div>Call center dédié + prise de RDV</div></li>
            <li><span class="dot"></span><div>Montage & suivi dossiers de financement (selon éligibilité)</div></li>
            <li><span class="dot"></span><div>Tunnel personnalisé + CRM + reporting mensuel</div></li>
            <li><span class="dot"></span><div>Interlocuteur unique + pilotage KPI</div></li>
          </ul>
        </div>

        <div class="section">
          <p class="note"><span class="ok">Inclus</span> : cadre RGPD de base (mentions + consentement). Hors : budget pub, achats d'outils tiers, frais d'hébergement spécifiques, démarches juridiques.</p>
        </div>
      </div>
    </div>

    <div class="grid3">
      <div class="card">
        <h2>Pour qui ?</h2>
        <p class="muted">TPE/PME, associations, centres de formation, professionnels de service qui veulent une organisation plus rentable et plus rigoureuse.</p>
      </div>
      <div class="card">
        <h2>Résultats attendus</h2>
        <p class="muted">Plus de RDV qualifiés, un suivi des demandes plus propre, un pipeline piloté, des dossiers financement structurés.</p>
      </div>
      <div class="card">
        <h2>Déploiement</h2>
        <p class="muted">On démarre par un diagnostic, puis un plan de mise en place (scripts, CRM, mails, tunnel, reporting).</p>
      </div>
    </div>

    <div class="card section" id="diagnostic">
      <h2>Demande de diagnostic</h2>
      <p class="muted">Copiez ce formulaire dans votre outil (WordPress, Typeform, Google Forms, Moodle). Le but est de qualifier vite et bien.</p>

      <form class="form" onsubmit="alert('Formulaire de démonstration. Connectez-le à votre outil (CRM / email / webhook).'); return false;">
        <div class="twocol">
          <div>
            <label>Nom / Prénom</label>
            <input required placeholder="Ex. Marie DUPONT"/>
          </div>
          <div>
            <label>Entreprise / Structure</label>
            <input required placeholder="Ex. ABC Formation"/>
          </div>
        </div>
        <div class="twocol">
          <div>
            <label>Email</label>
            <input type="email" required placeholder="exemple@entreprise.fr"/>
          </div>
          <div>
            <label>Téléphone</label>
            <input required placeholder="+590 ..."/>
          </div>
        </div>
        <div class="twocol">
          <div>
            <label>Type de structure</label>
            <select required>
              <option value="">Choisir…</option>
              <option>TPE / PME</option>
              <option>Centre de formation</option>
              <option>Association</option>
              <option>Profession libérale</option>
              <option>Autre</option>
            </select>
          </div>
          <div>
            <label>Budget mensuel envisagé</label>
            <select required>
              <option value="">Choisir…</option>
              <option>250€</option>
              <option>500€</option>
              <option>1 000€</option>
              <option>À discuter</option>
            </select>
          </div>
        </div>
        <div>
          <label>Besoins prioritaires (décrivez brièvement)</label>
          <textarea required placeholder="Ex. Gestion des appels, RDV, montage de financement, tunnel, CRM, relances..."></textarea>
        </div>
        <button class="btn primary" type="submit">Envoyer la demande</button>
        <p class="note">Astuce : connectez ce formulaire à votre CRM via webhook (Make/Zapier) ou via un plugin WordPress.</p>
      </form>
    </div>

    <div class="card section faq">
      <h2>Questions fréquentes</h2>
      <div class="faq" style="display:grid; gap:10px">
        <details>
          <summary>Est-ce 100% automatisé ?</summary>
          <p class="muted">Non. L'automatisation sert à gagner du temps (formulaires, RDV, emails, facturation). Les points critiques restent humains : qualification, montage financement, closing, suivi PREMIUM.</p>
        </details>
        <details>
          <summary>Quels financements pouvez-vous mobiliser ?</summary>
          <p class="muted">Selon l'éligibilité : dispositifs publics (France Travail, Région…), OPCO, fonds associatifs, etc. Nous structurons le dossier et faisons le suivi.</p>
        </details>
        <details>
          <summary>Combien de temps pour démarrer ?</summary>
          <p class="muted">Après validation : mise en place progressive (scripts, CRM, mails, tunnel) puis montée en charge.</p>
        </details>
      </div>
    </div>

    <footer>
      <div>© Votre structure — Tous droits réservés. Mentions légales & politique de confidentialité à ajouter.</div>
    </footer>
  </div>
</body>
</html>
"""
html_path = os.path.join(base, "page_vente_abonnement_pro.html")
with open(html_path, "w", encoding="utf-8") as f:
    f.write(html)

# 2) Call center scripts (TXT)
script_txt = """SCRIPTS CALL CENTER (INBOUND) — VERSION OPÉRATIONNELLE
==================================================

Objectif général
- Répondre vite, rassurer, qualifier et orienter (information / devis / financement / rendez-vous).
- Créer une trace CRM systématique (aucun appel sans fiche).
- Transformer : prise de rendez-vous ou demande de pièces / email de suivi.

A. OUVERTURE D'APPEL (20–30 secondes)
------------------------------------
Bonjour, vous êtes bien chez [NOM STRUCTURE]. [Prénom] à l'appareil. Comment puis-je vous aider ?

Si la personne est pressée :
Je comprends. En 30 secondes, je vous pose 3 questions pour vous orienter rapidement.

B. IDENTIFICATION (obligatoire)
------------------------------
1) Puis-je avoir votre nom et le nom de votre structure ?
2) Votre numéro de téléphone et votre email ?
3) Vous nous contactez pour : (1) information (2) devis (3) financement (4) site / emails (5) autre ?

C. QUALIFICATION RAPIDE (3 questions)
-------------------------------------
1) Votre besoin principal aujourd'hui, en une phrase ?
2) Votre délai : immédiat / 1 mois / 3 mois ?
3) Votre budget mensuel envisagé : 250 / 500 / 1000 / à discuter ?

D. ARBRE DE DÉCISION (routing)
------------------------------
1) DEMANDE D'INFORMATION
- Reformulation : Si je résume, vous souhaitez [X].
- Réponse courte + proposition : Je vous propose un diagnostic gratuit de 15–30 min.
- Action : prise de RDV (Calendly) + envoi email récapitulatif.

2) DEMANDE DE DEVIS
- Demander : type de structure, volume, besoin (call center ? financement ? tunnel ?).
- Action : "Je vous envoie un devis sous 24–48h ouvrées."
- Pièces : SIRET, adresse, contact facturation, description besoin, budget, urgence.

3) DEMANDE FINANCEMENT
- Qualification : statut (demandeur d'emploi / salarié / indépendant / association), dispositif visé si connu.
- Action : "Nous faisons un diagnostic éligibilité et une check-list des pièces."
- Pièces habituelles (à adapter) : pièce ID, RIB, justificatifs statut, devis/objectif, CV (si formation), etc.
- RDV : proposer un RDV avec le référent financement.

4) DEMANDE DIGITAL (domaine / mails / tunnel)
- Vérifier : avez-vous déjà un nom de domaine ? un site ? un logo ?
- Action : "Nous pouvons mettre en place nom de domaine + mails pro + page de vente + formulaire + base leads."
- RDV : proposer un RDV de cadrage (15–30 min).

E. PRISE DE RENDEZ-VOUS (script)
--------------------------------
Pour avancer, le plus efficace est un diagnostic. J'ai deux créneaux : [jour/heure] ou [jour/heure]. Quel créneau vous convient ?
- Confirmer : parfait. Je vous envoie une confirmation par email + SMS si possible.

F. GESTION D'UN CLIENT MÉCONTENT (script)
----------------------------------------
Je comprends votre frustration. Je vais m'occuper de votre demande.
- 1) Que s'est-il passé exactement ?
- 2) Quel résultat attendez-vous ?
- 3) Je vous propose : [solution] + délai [X].
- Toujours : créer ticket CRM + escalade si nécessaire.

G. CLÔTURE (obligatoire)
------------------------
Merci. Pour confirmer : je récapitule [action 1], [action 2], délai [X].
Vous pouvez nous recontacter au [téléphone] ou répondre à l'email de suivi.
Bonne journée.

H. CHAMPS CRM À REMPLIR (aucune exception)
------------------------------------------
- Date/heure, canal, motif
- Nom, entreprise, téléphone, email
- Segment (PRO / PREMIUM / START)
- Urgence, budget, besoin principal
- Prochaine action + responsable + échéance
- Statut (Nouveau / En cours / RDV pris / Devis envoyé / Gagné / Perdu)
"""
script_path = os.path.join(base, "scripts_call_center.txt")
with open(script_path, "w", encoding="utf-8") as f:
    f.write(script_txt)

# 3) Devis template (DOCX)
doc = Document()
styles = doc.styles
style = styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_paragraph("DEVIS — Abonnement Multi-Services (Call center + Financement + Digital)")
title.runs[0].bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

p = doc.add_paragraph("Fournisseur")
p.runs[0].bold = True
doc.add_paragraph("Raison sociale : [VOTRE RAISON SOCIALE]")
doc.add_paragraph("Adresse : [ADRESSE]")
doc.add_paragraph("SIRET : [SIRET]  |  TVA : [TVA]")
doc.add_paragraph("Email : [EMAIL]  |  Téléphone : [TEL]")

doc.add_paragraph("")
p = doc.add_paragraph("Client")
p.runs[0].bold = True
doc.add_paragraph("Raison sociale : [CLIENT]")
doc.add_paragraph("Adresse : [ADRESSE CLIENT]")
doc.add_paragraph("SIRET : [SIRET CLIENT]  |  TVA : [TVA CLIENT]")
doc.add_paragraph("Contact : [NOM]  |  Email : [EMAIL]  |  Téléphone : [TEL]")

doc.add_paragraph("")
meta = doc.add_paragraph("Références")
meta.runs[0].bold = True
doc.add_paragraph("N° Devis : [DEVIS-XXXX]")
doc.add_paragraph(f"Date : {datetime.date.today().strftime('%d/%m/%Y')}")
doc.add_paragraph("Validité du devis : 30 jours")
doc.add_paragraph("Périodicité : Mensuelle (abonnement)")
doc.add_paragraph("Démarrage : [DATE DE DÉMARRAGE]")

doc.add_paragraph("")
p = doc.add_paragraph("Objet")
p.runs[0].bold = True
doc.add_paragraph("Mise à disposition d'un service externalisé intégrant : centre d'appels, ingénierie de financement et infrastructure digitale (nom de domaine, mails pro, CRM et tunnel de vente).")

doc.add_paragraph("")
p = doc.add_paragraph("Détail des prestations")
p.runs[0].bold = True

table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text = "Réf."
hdr[1].text = "Prestation"
hdr[2].text = "Qté"
hdr[3].text = "PU HT (€)"
hdr[4].text = "Total HT (€)"

rows = [
    ("PRO", "Abonnement PRO (call center mutualisé + diagnostic financement + domaine/mails + tunnel standard + CRM)", "1", "[500]", "[500]"),
    ("PREM", "Abonnement PREMIUM (call center dédié + montage/suivi financement + tunnel personnalisé + CRM + reporting)", "1", "[1000]", "[1000]"),
    ("SETUP", "Frais de mise en place (paramétrage, scripts, CRM, emails, tunnel) — optionnel", "1", "[XXX]", "[XXX]"),
]
for r in rows:
    row_cells = table.add_row().cells
    for i, val in enumerate(r):
        row_cells[i].text = str(val)

doc.add_paragraph("")
tot = doc.add_paragraph("Récapitulatif")
tot.runs[0].bold = True
doc.add_paragraph("Total HT : [TOTAL_HT] €")
doc.add_paragraph("TVA : [TVA_%] % — [MONTANT_TVA] €")
doc.add_paragraph("Total TTC : [TOTAL_TTC] €")

doc.add_paragraph("")
p = doc.add_paragraph("Conditions")
p.runs[0].bold = True
doc.add_paragraph("• Paiement : mensuel, à l'avance, par prélèvement / carte / virement.")
doc.add_paragraph("• Résiliation : [préavis X jours] selon conditions générales.")
doc.add_paragraph("• Hors périmètre : budget publicitaire, achats d'outils tiers, démarches juridiques, contenus rédactionnels lourds, développement sur-mesure non prévu.")
doc.add_paragraph("• Données : le client reste responsable des données fournies ; un cadre RGPD de base est mis en place (mentions + consentement).")

doc.add_paragraph("")
sig = doc.add_paragraph("Acceptation du devis")
sig.runs[0].bold = True
doc.add_paragraph("Bon pour accord : ____________________________    Date : ____/____/______")
doc.add_paragraph("Nom / Qualité : _____________________________    Signature / Cachet : ____________________")

devis_path = os.path.join(base, "modele_devis_abonnement_multiservices.docx")
doc.save(devis_path)

# 4) CRM import template (CSV)
columns = [
    "LeadID","DateCreation","Source","Nom","Prenom","Entreprise","SIRET","Email","Telephone",
    "TypeStructure","BesoinPrincipal","OffreCible","BudgetMensuel","Urgence","StatutLead",
    "ProchaineAction","Responsable","Echeance","Notes"
]
df = pd.DataFrame([{
    "LeadID":"L-0001",
    "DateCreation":datetime.date.today().isoformat(),
    "Source":"Formulaire Web",
    "Nom":"DUPONT",
    "Prenom":"Marie",
    "Entreprise":"ABC Formation",
    "SIRET":"",
    "Email":"marie@abcformation.fr",
    "Telephone":"+590...",
    "TypeStructure":"Centre de formation",
    "BesoinPrincipal":"Call center + devis + financement",
    "OffreCible":"PREMIUM",
    "BudgetMensuel":"1000",
    "Urgence":"Immédiate",
    "StatutLead":"Nouveau",
    "ProchaineAction":"Prendre RDV diagnostic",
    "Responsable":"[Nom]",
    "Echeance":(datetime.date.today()+datetime.timedelta(days=2)).isoformat(),
    "Notes":"Lead test / exemple"
}], columns=columns)

crm_path = os.path.join(base, "template_import_crm_leads.csv")
df.to_csv(crm_path, index=False, encoding="utf-8")

# Module variables for external access
__all__ = ['html_path', 'script_path', 'devis_path', 'crm_path']
